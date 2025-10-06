"""Factories that materialise deterministic Office documents for tests."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook


def create_sample_commitments_workbook(destination: Path) -> Path:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Commitments"
    sheet.append(["Commitment ID", "Vendor", "Phase", "Amount"])
    sheet.append(["C-1001", "Acme Builders", "Design", 125000.50])
    sheet.append(["C-1002", "Zenith Construction", "Construction", 98000.00])
    sheet.append(["C-1003", "Acme Builders", "Inspection", 15750.25])
    workbook.save(destination)
    workbook.close()
    return destination


def create_sample_environment_doc(destination: Path) -> Path:
    document = Document()
    document.add_heading("Environment Overview", level=1)
    document.add_paragraph("Project: River Crossing")
    document.add_paragraph("Region: North")
    document.add_paragraph("Total Budget: $238,750.75")
    document.add_paragraph("Prepared By: Environmental Review Board")
    document.save(destination)
    return destination


__all__ = [
    "create_sample_commitments_workbook",
    "create_sample_environment_doc",
]
